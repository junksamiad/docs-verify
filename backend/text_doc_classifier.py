import os
import json
from docx import Document # For reading .docx files
from openai import OpenAI

client = OpenAI()

# --- Model Configuration ---
TEXT_CLASSIFIER_MODEL_ID = 'gpt-4o-2024-08-06' # Or a cheaper/faster model if sufficient
# --- End Model Configuration ---

def extract_text_from_docx(docx_path: str) -> str | None:
    """
    Extracts all text content from a .docx file.
    Returns the extracted text as a single string, or None on failure.
    """
    if not os.path.exists(docx_path):
        print(f"[TextDocClassifier ERROR] DOCX file not found at {docx_path}")
        return None
    try:
        document = Document(docx_path)
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        # Consider extracting text from tables as well if needed
        # for table in document.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             for para_in_cell in cell.paragraphs:
        #                 full_text.append(para_in_cell.text)
        extracted = '\n'.join(full_text)
        print(f"[TextDocClassifier LOG] Extracted {len(extracted)} characters from {docx_path}")
        return extracted
    except Exception as e:
        print(f"[TextDocClassifier ERROR] Failed to extract text from {docx_path}: {e}")
        return None

def classify_text_document_type(text_content: str, document_types: list[str]) -> dict | None:
    """
    Classifies the type of document based on its text content using an LLM.
    Returns a dictionary containing the predicted type and the original text, or None.
    """
    if not text_content:
        print("[TextDocClassifier ERROR] Text content cannot be empty for classification.")
        return None
    if not document_types:
        print("[TextDocClassifier ERROR] Document types list cannot be empty.")
        return None
    
    if not os.environ.get("OPENAI_API_KEY"):
        print("[TextDocClassifier ERROR] OPENAI_API_KEY environment variable not set.")
        return None

    output_schema = {
        "type": "object",
        "properties": {
            "predicted_document_type": {
                "type": "string",
                "enum": document_types,
                "description": "The classified type of the document from the provided list based on its text content."
            }
        },
        "required": ["predicted_document_type"]
    }

    prompt = (
        f"Please analyze the following text extracted from a document and classify its type. "
        f"The document type must be one of the following: {', '.join(document_types)}. "
        f"Respond with a JSON object that strictly adheres to the following schema: {json.dumps(output_schema)}. "
        f"Ensure the 'predicted_document_type' field contains only one of the allowed types.\n\n"
        f"Document Text:\n---BEGIN TEXT---\n{text_content[:15000]} \n---END TEXT---"
    ) # Truncate text if very long to manage token limits, adjust as needed

    api_response_text = None
    try:
        # Corrected f-string for logging, ensuring quotes are balanced
        log_text_snippet = text_content[:200].replace('\n', ' ')
        print(f"[TextDocClassifier LOG] Sending text (first 200 chars: '{log_text_snippet}') to {TEXT_CLASSIFIER_MODEL_ID} for classification.")
        response = client.chat.completions.create(
            model=TEXT_CLASSIFIER_MODEL_ID,
            messages=[
                {"role": "system", "content": "You are an expert document classifier. Your task is to determine the document type based on the provided text and a list of allowed types. Respond in JSON format according to the schema provided."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )

        if response.choices and response.choices[0].message and response.choices[0].message.content:
            api_response_text = response.choices[0].message.content
            print(f"[TextDocClassifier LOG] Raw JSON response: {api_response_text}")
            parsed_json = json.loads(api_response_text)
            predicted_type = parsed_json.get("predicted_document_type")

            if predicted_type and predicted_type in document_types:
                return {"predicted_document_type": predicted_type, "extracted_text": text_content}
            elif predicted_type:
                print(f"[TextDocClassifier WARNING] Model returned type '{predicted_type}' not in allowed list: {document_types}. Raw JSON: {api_response_text}")
                return {"predicted_document_type": "Other", "extracted_text": text_content} # Fallback or error
            else:
                print(f"[TextDocClassifier ERROR] 'predicted_document_type' missing in JSON: {api_response_text}")
                return None
        else:
            print(f"[TextDocClassifier ERROR] No valid content in API response. Full response: {response}")
            return None

    except Exception as e:
        print(f"[TextDocClassifier ERROR] API call/processing failed: {e}. API response text: {api_response_text if api_response_text else 'N/A'}")
        return None

if __name__ == '__main__':
    # Create a dummy .docx file for testing
    if not os.path.exists("temp_test_files"): os.makedirs("temp_test_files")
    dummy_docx_path = os.path.join("temp_test_files", "dummy_cv.docx")
    doc = Document()
    doc.add_heading('John Doe - Curriculum Vitae', 0)
    doc.add_paragraph('Contact: john.doe@email.com | (555) 123-4567')
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Highly motivated software engineer with 5 years of experience...')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Engineer at Tech Solutions Inc. (2019-Present)')
    doc.add_paragraph('- Developed and maintained web applications.')
    doc.add_heading('Education', level=1)
    doc.add_paragraph('BSc Computer Science - University of Example (2019)')
    doc.save(dummy_docx_path)

    dummy_letter_path = os.path.join("temp_test_files", "dummy_letter.docx")
    doc_letter = Document()
    doc_letter.add_paragraph('Dear Sir/Madam,\n\nI am writing to apply for the position...\n\nSincerely,\nJane Smith')
    doc_letter.save(dummy_letter_path)

    possible_types = ["CV", "Letter", "Report", "Invoice", "Other"]

    print(f"\n--- Test 1: CV Classification from DOCX ---")
    extracted_text_cv = extract_text_from_docx(dummy_docx_path)
    if extracted_text_cv:
        classification_cv = classify_text_document_type(extracted_text_cv, possible_types)
        print(f"CV Classification Result: {classification_cv.get('predicted_document_type') if classification_cv else 'Failed'}")

    print(f"\n--- Test 2: Letter Classification from DOCX ---")
    extracted_text_letter = extract_text_from_docx(dummy_letter_path)
    if extracted_text_letter:
        classification_letter = classify_text_document_type(extracted_text_letter, possible_types)
        print(f"Letter Classification Result: {classification_letter.get('predicted_document_type') if classification_letter else 'Failed'}")

    # Clean up dummy files
    # os.remove(dummy_docx_path)
    # os.remove(dummy_letter_path)
    # shutil.rmtree("temp_test_files")
    print("\nNote: Dummy files created in temp_test_files. Clean up manually if needed.") 